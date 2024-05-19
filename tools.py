"""
作者： 河南理工大学 曹显嵩
    河南理工大学 吴锡煐
 因为算力设备的问题，本项目仅仅使用了百度千帆模型
    只需要更改ai.deal()函数的设置便可以进行自定义其他的大语言模型以及本地部署的模型
    本项目提供了prompt自定义智能体，通过配置config.json
    以及文本转语音
    语音转文本请自行完成，文字内容直接传入ai.deal()便可

    若各位是有设备条件进行微调可以选择忽视自定义智能体，若大家如我一样缺乏设备，可选择自定义智能体来完成角色的设定，毕竟目前领先的各种大模型已经是适用多种场景。
    自定义prompt是考虑到使用的方便性，以及对算力的考虑。
    本项目，主要功能是基于api实现，文本转语音模型可以部署到自己电脑进行内网穿透映射到公网进行发送请求，能够实现各种显存低下设备的部署，对算力要求也不高
    综合考虑，国赛设备显存远远不够，又无其他算力设备支持，这是性价比最高的方案


缺陷:
    pdf文件上传暂时没做(时间关系):
        实现思路：
            1.api
            2.ocr识别->调用翻译智能体->写入
                orc项目推荐 paddleocr
                        chineseocr_lite

项目用途：仅供参考。

"""

import requests
import urllib
from pydub import AudioSegment
import requests
from pydub.playback import play
import json
import docx
import os
from pdf2docx import Converter
from docx2pdf import convert

#百度千帆大模型的key
#https://qianfan.cloud.baidu.com/
#也可去调用chatgpt,具体调用方法我便不写了，也相信能走到国赛的同志，知道如何去改

API_KEY  = ''
SECRET_KEY = ''
config_path = './config.json'

class ai():


    def __init__(self):
        #导入自定义的prompt智能体
        self.translater_object = translater()

        self.history = {"messages": [],
        "temperature": 0.8, #较高的数值会使输出更加随机，而较低的数值会使其更加集中和确定
        "top_p": 0.8,       #影响输出文本的多样性，取值越大，生成文本的多样性越强
        "penalty_score": 1,
        "disable_search": False,
        "enable_citation": False,
        "response_format": "text"
    }

        self.headers = {
        'Content-Type': 'application/json'
    }

        self.API_KEY = API_KEY
        self.SECRET_KEY = SECRET_KEY
        self.url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie-3.5-8k-0205?access_token=" + self._get_access_token()

    def _get_access_token(self):
        """
        使用 AK，SK 生成鉴权签名（Access Token）
        :return: access_token，或是None(如果错误)
        """
        url = "https://aip.baidubce.com/oauth/2.0/token"
        params = {"grant_type": "client_credentials", "client_id": self.API_KEY, "client_secret": self.SECRET_KEY}
        return str(requests.post(url, params=params).json().get("access_token"))

    def _define_messg(self,Content,role ='user'):
        temp = {}
        temp["content"] = Content
        temp["role"] = role

        self.history['messages'].append(temp)
    # def _pdf_to_word(self,pdf_path):
    #     """
    #
    #     :param pdf_path:pdf 路径
    #     :param word_path: word保存路径
    #     :return: 转换成为word 路径
    #     """
    #     cv = Converter(pdf_path)
    #     word_path = pdf_path.replace('.pdf','.docx')
    #     cv.convert(word_path, start=0, end=None)
    #     cv.close()
    #     return word_path
    # def _docx_to_pdf(self,docx_path):
    #     """
    #
    #     :param docx_path:word文档路径
    #     :param pdf_path: 保存的pdf路径
    #     :return: 转换成为pdf文档 路径
    #     """
    #     pdf_path=docx_path.replace('.docx','.pdf')
    #     convert(docx_path,pdf_path)
    #
    #     return pdf_path
    #
    # def _pdf_to_translate(self, pdf_path):
    #     """
    #
    #     :param pdf_path: pdf 路径
    #     :return: path
    #     """
    #     word = self._pdf_to_word(pdf_path)
    #     word_translate = self.file_to_file(word)
    #
    #     return self._docx_to_pdf(word_translate)
    #

        

    def _word_to_text(self,word_path):
        """

        :param word_path: word_路径
        :return:file path
        """
        # 获取文档对象
        file = docx.Document(word_path)

        # 输出每一段的内容
        for para in file.paragraphs:
            temp = para.text
            if temp == "":
                continue
            temp = self.translater_object.translate(temp) #翻译过的结果

            para.text = temp

        new_path = f"{word_path.split('.')[0]}_translated.docx"

        file.save(new_path)

        return os.path.join(os.getcwd(),new_path)

    def file_to_file(self,file_path):
        """

        :param file_path: 传入文件路径
        :return: 返回保存文件地址
        """
        if file_path.split('.')[-1] == 'pdf':
            pass
            # return self._pdf_to_translate(file_path)
        elif file_path.split('.')[-1] == 'doc' or file_path.split('.')[-1] == 'docx':
            return self._word_to_text(file_path)


    def Summary_to_text(self,file_path):
        """
        读取文档所有内容，进行deal模块，得到结果

        :param file_path: 文本路径
        :return: str总结文本
        """
        text_ls = []
        if file_path.split('.')[-1] == 'pdf':
            pass

        elif file_path.split('.')[-1] == 'doc' or file_path.split('.')[-1] == 'docx':
            file = docx.Document(file_path)

            # 输出每一段的内容
            for para in file.paragraphs:
                text_ls.append(para.text)

            text = '\n'.join(text_ls)
            return self.deal('请对以下进行总结\n'+text)


    def deal(self,content):
        """
        在这个函数中，只需要更换gpt对话的模型，便可正常运行



        :param content:传入文本
        :return: 返回ai结果
        """
        self._define_messg(content)
        payload = json.dumps(self.history)


        response = requests.request('POST',self.url,headers=self.headers,data=payload)
        output = response.json()['result']
        self._define_messg(Content=output,role='assistant')



        return output

    def text_to_speech(self,text):
        """
        采用了https://www.yuque.com/baicaigongchang1145haoyuangong/ib3g1e 项目
        具体详情请关注；
        使用前请先启动 GPT-SoVITS-Inference/0 一键启动脚本/5 启动后端程序.bat
        或是启动 GPT-SoVITS-Inference/Inference/src/tts_backend.py
        此文本转语音项目可以部署在本地进行内网穿透，或是部署在gpu服务器:
        推荐几家：
            https://gpushare.com/
            https://www.ucloud.cn/
            ucloud仅仅考虑新客优惠
        :param text:文本（str）
        :return: 返回转换音频文件的缓存地址
        """
        url = urllib.parse.quote(text)
        #此为文本转语音的请求地址与发生get请求的参数
        wav = requests.get(
            f'http://127.0.0.1:5000/tts?character=胡桃&text={url}')

        wav = wav.content
        with open('./temp.wav', 'wb') as fp:
            fp.write(wav)

        # re = AudioSegment.from_wav('temp.wav')
        # play(re)
        return './temp.wav'


class translater(ai):
    """
    本设置是为了自定义一个智能体
    继承自ai类


    """
    def __init__(self):

        self.history = {"messages": [],
                        "temperature": 0.8,
                        "top_p": 0.8,
                        "penalty_score": 1,
                        "disable_search": False,
                        "enable_citation": False,
                        "response_format": "text"
                        }

        self.headers = {
            'Content-Type': 'application/json'
        }

        self.API_KEY = API_KEY
        self.SECRET_KEY = SECRET_KEY
        self.url = "https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie-3.5-8k-0205?access_token=" + self._get_access_token()
        self.config_path = config_path
        self._get_config()
    def _get_config(self):
        """
        读取预设置的 prompt 指令

        """
        with open(self.config_path,'r',encoding='utf8') as f:
            self.config = json.load(f)



    def _define_translate(self,content,role='翻译'):
        """
        采取格式化输入，保证每次输出符合预设目标


        : param role:自定义智能体角色名称，使用json文件存储
        : param content: prompt指令，进行自定义智能体，不局限于翻译
        :return : 返回封加前缀的 prompt 指令
        """
        prompt = self.config.get(role)

        if prompt:
            return '请返回无此角色'


        return str(prompt).format(content)



    def translate(self,content):
        """

        :param content: 传入文本
        :return: 返回翻译的值
        """

        content = self._define_translate(content)
        return self.deal(content)





if __name__ =='__main__':

    app = ai()

    while True:
        t = str(input("请输入问题:"))
        if t[0]=='/':
            t = t[1:]

            aw = app.translater_object.translate(t)


        else:
            aw = app.deal(t)

        app.text_to_speech(aw)
        print(aw)

